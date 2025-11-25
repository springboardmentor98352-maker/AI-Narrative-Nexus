import os
import csv
from typing import List

try:
    from docx import Document
    DOCX_AVAILABLE = True
except Exception:
    DOCX_AVAILABLE = False


def ensure_data_dir(data_dir: str = "sample_data") -> str:
    """Ensure the data directory exists and return its path."""
    if not os.path.exists(data_dir):
        os.makedirs(data_dir)
    return data_dir


class _DummyUploaded:
    """Internal helper used for tests or compatibility if needed."""

    def __init__(self, name: str, data: bytes):
        self.name = name
        self._data = data

    def getbuffer(self) -> bytes:
        return self._data


def save_uploaded_file(uploaded_file, data_dir: str = "sample_data") -> str:
    """Save a Streamlit uploaded file-like object to disk and return path."""
    ensure_data_dir(data_dir)
    file_path = os.path.join(data_dir, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    return file_path


def get_sample_files(data_dir: str = "sample_data") -> List[str]:
    ensure_data_dir(data_dir)
    return sorted(os.listdir(data_dir))


def parse_preview(file_path: str, max_lines: int = 20) -> str:
    """Return a text preview for supported file types (.txt, .csv, .docx).

    - For .txt: return first `max_lines` lines.
    - For .csv: return first `max_lines` rows as comma-separated strings.
    - For .docx: return first `max_lines` paragraphs (if python-docx installed).
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                lines = []
                for i, line in enumerate(f):
                    if i >= max_lines:
                        break
                    lines.append(line.rstrip("\n"))
            return "\n".join(lines)
        except Exception as e:
            return f"[Error reading text file: {e}]"

    if ext == ".csv":
        try:
            with open(file_path, newline="", encoding="utf-8", errors="replace") as csvfile:
                reader = csv.reader(csvfile)
                lines = []
                for i, row in enumerate(reader):
                    if i >= max_lines:
                        break
                    # join with comma+space for nicer preview
                    lines.append(", ".join(row))
            return "\n".join(lines)
        except Exception as e:
            return f"[Error reading CSV file: {e}]"

    if ext == ".docx":
        if not DOCX_AVAILABLE:
            return "[python-docx not installed — cannot preview .docx files]"
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs[:max_lines]]
            return "\n".join(paragraphs)
        except Exception as e:
            return f"[Error reading DOCX file: {e}]"

    return f"[Preview not available for {ext} files]"


def read_full_text(file_path: str) -> str:
    """Read and return the full text content from a file.
    
    Supported formats: .txt, .csv, .docx
    - For .txt: returns all text content
    - For .csv: intelligently extracts text columns (tries 'text', 'content', 'message' columns first)
    - For .docx: extracts all paragraph text
    
    This function performs basic cleaning:
    - Replaces multiple whitespace/newlines with single spaces
    - Strips leading/trailing whitespace
    - Handles encoding errors gracefully
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()

    if ext == ".txt":
        try:
            with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                text = f.read()
            # Basic cleaning: normalize whitespace
            text = " ".join(text.split())
            return text
        except Exception as e:
            return f"[Error reading text file: {e}]"

    if ext == ".csv":
        try:
            with open(file_path, newline="", encoding="utf-8", errors="replace") as csvfile:
                reader = csv.DictReader(csvfile)
                all_text = []
                
                # Try to find text-like columns
                first_row = None
                text_columns = []
                
                for row in reader:
                    if first_row is None:
                        first_row = row
                        # Look for common text column names
                        for col in ['text', 'content', 'message', 'body', 'description', 'tweet', 'post', 'comment']:
                            if col in row:
                                text_columns.append(col)
                        
                        # If no common columns found, use all columns
                        if not text_columns:
                            text_columns = list(row.keys())
                    
                    # Extract text from identified columns
                    row_text = " ".join([str(row.get(col, "")) for col in text_columns])
                    if row_text.strip():
                        all_text.append(row_text)
                
            # Join all rows with space and normalize whitespace
            text = " ".join(all_text)
            text = " ".join(text.split())
            return text
        except Exception as e:
            return f"[Error reading CSV file: {e}]"

    if ext == ".docx":
        if not DOCX_AVAILABLE:
            return "[python-docx not installed — cannot read .docx files]"
        try:
            doc = Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs]
            text = " ".join(paragraphs)
            # Normalize whitespace
            text = " ".join(text.split())
            return text
        except Exception as e:
            return f"[Error reading DOCX file: {e}]"

    return f"[Cannot read {ext} files — unsupported format]"


def get_text_summary(text: str, max_words: int = 100) -> str:
    """Generate an intelligent summary by analyzing the text content.
    
    Args:
        text: The text to summarize
        max_words: Maximum number of words to consider for analysis
    
    Returns:
        A descriptive summary of the text content
    """
    if not text or len(text.strip()) == 0:
        return "Empty text file with no content."
    
    words = text.split()
    if len(words) == 0:
        return "No readable content found in the file."
    
    # Analyze content characteristics
    total_words = len(words)
    unique_words = len(set([w.lower() for w in words]))
    
    # Sample text for analysis (first 1000 words)
    sample_text = " ".join(words[:1000]).lower()
    
    # Detect content type based on keywords
    content_type = "general text"
    context_info = []
    
    # Check for social media content
    if "@" in sample_text or "tweet" in sample_text or "retweet" in sample_text:
        content_type = "social media posts"
        if "airline" in sample_text or "flight" in sample_text:
            context_info.append("about airlines and flight experiences")
        if "sentiment" in sample_text or "positive" in sample_text or "negative" in sample_text:
            context_info.append("with sentiment analysis data")
    
    # Check for reviews or feedback
    elif any(word in sample_text for word in ["review", "rating", "customer", "feedback"]):
        content_type = "customer reviews or feedback"
    
    # Check for news or articles
    elif any(word in sample_text for word in ["article", "report", "news", "journalist"]):
        content_type = "news articles or reports"
    
    # Check for conversational data
    elif any(word in sample_text for word in ["chat", "message", "conversation", "reply"]):
        content_type = "conversational messages"
    
    # Check for product/service mentions
    if any(brand in sample_text for brand in ["virgin", "united", "delta", "american airlines", "southwest"]):
        if not context_info:
            context_info.append("related to airline services")
    
    # Build summary
    summary_parts = [f"This file contains {content_type}"]
    
    if context_info:
        summary_parts.append(" " + ", ".join(context_info))
    
    summary_parts.append(f". It has approximately {total_words:,} words with {unique_words:,} unique terms")
    
    # Add sentiment/tone detection
    sentiment_words = {
        "positive": ["good", "great", "excellent", "love", "best", "happy", "thanks", "amazing"],
        "negative": ["bad", "worst", "terrible", "hate", "awful", "poor", "disappointed", "angry"],
    }
    
    pos_count = sum(1 for word in sentiment_words["positive"] if word in sample_text)
    neg_count = sum(1 for word in sentiment_words["negative"] if word in sample_text)
    
    if pos_count > neg_count * 1.5:
        summary_parts.append(", predominantly positive in tone")
    elif neg_count > pos_count * 1.5:
        summary_parts.append(", predominantly negative in tone")
    elif pos_count > 0 and neg_count > 0:
        summary_parts.append(", with mixed sentiments")
    
    summary_parts.append(".")
    
    return "".join(summary_parts)
